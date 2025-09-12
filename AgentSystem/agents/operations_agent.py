
"""
Operations Automation Agent - AgentSystem Profit Machine
Handles document processing, workflow automation, and business process optimization
"""

import asyncio
import json
import logging
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union, Tuple
from enum import Enum
from dataclasses import dataclass, asdict
from pathlib import Path
import asyncpg
import aiofiles
import aiohttp
from pydantic import BaseModel, Field
import pandas as pd
import numpy as np
from io import BytesIO
import base64
import hashlib
import uuid

# Document processing imports
import PyPDF2
import docx
from PIL import Image
import pytesseract
import openpyxl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# AI processing imports
import openai
from anthropic import Anthropic

logger = logging.getLogger(__name__)

class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    XLSX = "xlsx"
    IMAGE = "image"
    EMAIL = "email"
    JSON = "json"
    XML = "xml"
    HTML = "html"

class ProcessingStatus(Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    QUEUED = "queued"

class WorkflowStatus(Enum):
    DRAFT = "draft"
    ACTIVE = "active"
    PAUSED = "paused"
    COMPLETED = "completed"
    FAILED = "failed"

class AutomationType(Enum):
    DOCUMENT_EXTRACTION = "document_extraction"
    DATA_TRANSFORMATION = "data_transformation"
    EMAIL_PROCESSING = "email_processing"
    REPORT_GENERATION = "report_generation"
    APPROVAL_WORKFLOW = "approval_workflow"
    DATA_VALIDATION = "data_validation"
    NOTIFICATION_SYSTEM = "notification_system"
    BATCH_PROCESSING = "batch_processing"

@dataclass
class DocumentMetadata:
    document_id: str
    filename: str
    document_type: DocumentType
    file_size: int
    upload_timestamp: datetime
    tenant_id: str
    user_id: str
    checksum: str
    mime_type: str
    processing_status: ProcessingStatus
    extracted_text: Optional[str] = None
    metadata: Dict[str, Any] = None
    tags: List[str] = None
    confidence_score: float = 0.0

@dataclass
class ProcessingResult:
    document_id: str
    processing_type: str
    status: ProcessingStatus
    extracted_data: Dict[str, Any]
    confidence_score: float
    processing_time_seconds: float
    error_message: Optional[str] = None
    suggestions: List[str] = None

@dataclass
class WorkflowDefinition:
    workflow_id: str
    name: str
    description: str
    tenant_id: str
    automation_type: AutomationType
    trigger_conditions: Dict[str, Any]
    processing_steps: List[Dict[str, Any]]
    output_configuration: Dict[str, Any]
    status: WorkflowStatus
    created_by: str
    created_at: datetime
    last_modified: datetime
    execution_count: int = 0
    success_rate: float = 0.0

@dataclass
class WorkflowExecution:
    execution_id: str
    workflow_id: str
    tenant_id: str
    status: ProcessingStatus
    input_data: Dict[str, Any]
    output_data: Optional[Dict[str, Any]]
    started_at: datetime
    completed_at: Optional[datetime]
    error_message: Optional[str]
    processing_logs: List[str]
    execution_time_seconds: Optional[float] = None

class OperationsAgent:
    """
    Advanced Operations Automation Agent for document processing and workflow automation
    """

    def __init__(self, tenant_id: str, db_pool: asyncpg.Pool,
                 openai_client: openai.AsyncOpenAI, anthropic_client: Anthropic,
                 storage_path: str = "./document_storage"):
        self.tenant_id = tenant_id
        self.db_pool = db_pool
        self.openai_client = openai_client
        self.anthropic_client = anthropic_client
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)

        # Processing capabilities
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_txt,
            'csv': self._process_csv,
            'xlsx': self._process_xlsx,
            'image': self._process_image,
            'email': self._process_email,
            'json': self._process_json
        }

        # Workflow processors
        self.workflow_processors = {
            AutomationType.DOCUMENT_EXTRACTION: self._execute_document_extraction,
            AutomationType.DATA_TRANSFORMATION: self._execute_data_transformation,
            AutomationType.EMAIL_PROCESSING: self._execute_email_processing,
            AutomationType.REPORT_GENERATION: self._execute_report_generation,
            AutomationType.APPROVAL_WORKFLOW: self._execute_approval_workflow,
            AutomationType.DATA_VALIDATION: self._execute_data_validation,
            AutomationType.NOTIFICATION_SYSTEM: self._execute_notification_system,
            AutomationType.BATCH_PROCESSING: self._execute_batch_processing
        }

    async def upload_document(self, file_content: bytes, filename: str,
                            user_id: str, metadata: Optional[Dict] = None) -> DocumentMetadata:
        """Upload and process a document"""

        # Generate document ID and checksum
        document_id = str(uuid.uuid4())
        checksum = hashlib.sha256(file_content).hexdigest()

        # Determine document type
        file_extension = Path(filename).suffix.lower().lstrip('.')
        document_type = DocumentType(file_extension) if file_extension in [dt.value for dt in DocumentType] else DocumentType.TXT

        # Store file
        file_path = self.storage_path / self.tenant_id / document_id
        file_path.parent.mkdir(parents=True, exist_ok=True)

        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)

        # Create document metadata
        doc_metadata = DocumentMetadata(
            document_id=document_id,
            filename=filename,
            document_type=document_type,
            file_size=len(file_content),
            upload_timestamp=datetime.now(),
            tenant_id=self.tenant_id,
            user_id=user_id,
            checksum=checksum,
            mime_type=self._get_mime_type(filename),
            processing_status=ProcessingStatus.PENDING,
            metadata=metadata or {},
            tags=[]
        )

        # Store metadata in database
        await self._store_document_metadata(doc_metadata)

        # Queue for processing
        await self._queue_document_processing(document_id)

        return doc_metadata

    async def process_document(self, document_id: str,
                             processing_options: Optional[Dict] = None) -> ProcessingResult:
        """Process a document and extract structured data"""

        start_time = datetime.now()

        try:
            # Get document metadata
            doc_metadata = await self._get_document_metadata(document_id)
            if not doc_metadata:
                raise ValueError(f"Document {document_id} not found")

            # Update status
            await self._update_processing_status(document_id, ProcessingStatus.PROCESSING)

            # Load document content
            file_path = self.storage_path / self.tenant_id / document_id
            async with aiofiles.open(file_path, 'rb') as f:
                file_content = await f.read()

            # Process based on document type
            processor = self.supported_formats.get(doc_metadata.document_type.value)
            if not processor:
                raise ValueError(f"Unsupported document type: {doc_metadata.document_type.value}")

            extracted_data = await processor(file_content, processing_options or {})

            # Enhance with AI analysis
            ai_analysis = await self._ai_analyze_document(extracted_data, doc_metadata)
            extracted_data.update(ai_analysis)

            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()

            # Create result
            result = ProcessingResult(
                document_id=document_id,
                processing_type=doc_metadata.document_type.value,
                status=ProcessingStatus.COMPLETED,
                extracted_data=extracted_data,
                confidence_score=extracted_data.get('confidence_score', 0.8),
                processing_time_seconds=processing_time
            )

            # Update database
            await self._update_processing_status(document_id, ProcessingStatus.COMPLETED)
            await self._store_processing_result(result)

            return result

        except Exception as e:
            logger.error(f"Document processing failed for {document_id}: {str(e)}")

            processing_time = (datetime.now() - start_time).total_seconds()

            result = ProcessingResult(
                document_id=document_id,
                processing_type="unknown",
                status=ProcessingStatus.FAILED,
                extracted_data={},
                confidence_score=0.0,
                processing_time_seconds=processing_time,
                error_message=str(e)
            )

            await self._update_processing_status(document_id, ProcessingStatus.FAILED)
            await self._store_processing_result(result)

            return result

    async def _process_pdf(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process PDF document"""

        extracted_data = {
            'text_content': '',
            'page_count': 0,
            'metadata': {},
            'tables': [],
            'images': []
        }

        try:
            # Extract text using PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            extracted_data['page_count'] = len(pdf_reader.pages)

            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text_content.append(page_text)

            extracted_data['text_content'] = '\n'.join(text_content)
            extracted_data['metadata'] = pdf_reader.metadata or {}

            # Enhanced extraction with AI if needed
            if options.get('ai_enhancement', True):
                ai_extracted = await self._ai_extract_pdf_structure(extracted_data['text_content'])
                extracted_data.update(ai_extracted)

        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            extracted_data['error'] = str(e)

        return extracted_data

    async def _process_docx(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process DOCX document"""

        extracted_data = {
            'text_content': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {}
        }

        try:
            doc = docx.Document(BytesIO(file_content))

            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            extracted_data['paragraphs'] = paragraphs
            extracted_data['text_content'] = '\n'.join(paragraphs)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            extracted_data['tables'] = tables

            # Document properties
            extracted_data['metadata'] = {
                'author': doc.core_properties.author,
                'title': doc.core_properties.title,
                'subject': doc.core_properties.subject,
                'created': str(doc.core_properties.created),
                'modified': str(doc.core_properties.modified)
            }

        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            extracted_data['error'] = str(e)

        return extracted_data

    async def _process_txt(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process text document"""

        try:
            text_content = file_content.decode('utf-8')

            extracted_data = {
                'text_content': text_content,
                'line_count': len(text_content.split('\n')),
                'word_count': len(text_content.split()),
                'character_count': len(text_content)
            }

            # AI analysis for structure detection
            if options.get('ai_analysis', True) and len(text_content) > 100:
                ai_analysis = await self._ai_analyze_text_structure(text_content)
                extracted_data.update(ai_analysis)

            return extracted_data

        except Exception as e:
            logger.error(f"Text processing error: {str(e)}")
            return {'error': str(e), 'text_content': ''}

    async def _process_csv(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process CSV document"""

        try:
            # Read CSV data
            df = pd.read_csv(BytesIO(file_content))

            extracted_data = {
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': df.columns.tolist(),
                'data_types': df.dtypes.to_dict(),
                'sample_data': df.head(5).to_dict('records'),
                'statistics': {}
            }

            # Generate statistics for numeric columns
            numeric_columns = df.select_dtypes(include=[np.number]).columns
            if len(numeric_columns) > 0:
                extracted_data['statistics'] = df[numeric_columns].describe().to_dict()

            # AI-powered data analysis
            if options.get('ai_analysis', True):
                ai_insights = await self._ai_analyze_tabular_data(df)
                extracted_data.update(ai_insights)

            return extracted_data

        except Exception as e:
            logger.error(f"CSV processing error: {str(e)}")
            return {'error': str(e), 'row_count': 0, 'column_count': 0}

    async def _process_xlsx(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process Excel document"""

        try:
            # Read Excel file
            workbook = openpyxl.load_workbook(BytesIO(file_content))

            extracted_data = {
                'sheet_names': workbook.sheetnames,
                'sheets': {},
                'total_cells': 0
            }

            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Convert to DataFrame for analysis
                data = []
                for row in sheet.iter_rows(values_only=True):
                    data.append(row)

                if data:
                    df = pd.DataFrame(data[1:], columns=data[0])  # First row as headers

                    sheet_data = {
                        'row_count': len(df),
                        'column_count': len(df.columns),
                        'columns': df.columns.tolist(),
                        'sample_data': df.head(3).to_dict('records')
                    }

                    extracted_data['sheets'][sheet_name] = sheet_data
                    extracted_data['total_cells'] += len(df) * len(df.columns)

            return extracted_data

        except Exception as e:
            logger.error(f"Excel processing error: {str(e)}")
            return {'error': str(e), 'sheet_names': []}

    async def _process_image(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process image document with OCR"""

        try:
            # Open image
            image = Image.open(BytesIO(file_content))

            extracted_data = {
                'image_size': image.size,
                'image_mode': image.mode,
                'format': image.format,
                'text_content': '',
                'confidence_score': 0.0
            }

            # OCR text extraction
            if options.get('ocr_enabled', True):
                try:
                    ocr_text = pytesseract.image_to_string(image)
                    extracted_data['text_content'] = ocr_text

                    # Get confidence scores
                    ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                    confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                    if confidences:
                        extracted_data['confidence_score'] = sum(confidences) / len(confidences) / 100.0

                except Exception as ocr_error:
                    logger.warning(f"OCR processing failed: {str(ocr_error)}")
                    extracted_data['ocr_error'] = str(ocr_error)

            return extracted_data

        except Exception as e:
            logger.error(f"Image processing error: {str(e)}")
            return {'error': str(e), 'image_size': (0, 0)}

    async def _process_email(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process email document"""

        try:
            email_text = file_content.decode('utf-8')

            # Basic email parsing (would use email library for .eml files)
            extracted_data = {
                'content': email_text,
                'headers': {},
                'body': '',
                'attachments': []
            }

            # AI-powered email analysis
            if options.get('ai_analysis', True):
                ai_analysis = await self._ai_analyze_email(email_text)
                extracted_data.update(ai_analysis)

            return extracted_data

        except Exception as e:
            logger.error(f"Email processing error: {str(e)}")
            return {'error': str(e), 'content': ''}

    async def _process_json(self, file_content: bytes, options: Dict) -> Dict[str, Any]:
        """Process JSON document"""

        try:
            json_data = json.loads(file_content.decode('utf-8'))

            extracted_data = {
                'json_data': json_data,
                'structure_analysis': self._analyze_json_structure(json_data),
                'data_types': self._get_json_data_types(json_data)
            }

            return extracted_data

        except Exception as e:
            logger.error(f"JSON processing error: {str(e)}")
            return {'error': str(e), 'json_data': {}}

    async def _ai_analyze_document(self, extracted_data: Dict, metadata: DocumentMetadata) -> Dict[str, Any]:
        """AI-powered document analysis"""

        try:
            # Prepare content for AI analysis
            content = extracted_data.get('text_content', '')
            if not content:
                return {'ai_analysis': 'No text content available for analysis'}

            # Use OpenAI for document analysis
            response = await self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert document analyzer. Analyze the document and extract key insights, entities, and structure."
                    },
                    {
                        "role": "user",
                        "content": f"Analyze this {metadata.document_type.value} document:\n\n{content[:4000]}"
                    }
                ],
                max_tokens=1000
            )

            analysis = response.choices[0].message.content

            return {
                'ai_analysis': analysis,
                'key_entities': await self._extract_entities(content),
                'document_classification': await self._classify_document(content),
                'confidence_score': 0.85
            }

        except Exception as e:
            logger.error(f"AI analysis error: {str(e)}")
            return {'ai_analysis_error': str(e)}

    async def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities from text"""

        try:
            # Simple entity extraction (would use spaCy or similar in production)
            entities = []

            # Use AI for entity extraction
            response = await self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": "Extract named entities (people, organizations, locations, dates, amounts) from the text. Return as JSON array."
                    },
                    {
                        "role": "user",
                        "content": text[:2000]
                    }
                ],
                max_tokens=500
            )

            entities_text = response.choices[0].message.content
            try:
                entities = json.loads(entities_text)
            except:
                entities = [{'type': 'analysis', 'value': entities_text}]

            return entities

        except Exception as e:
            logger.error(f"Entity extraction error: {str(e)}")
            return []

    async def create_workflow(self, workflow_definition: Dict[str, Any]) -> WorkflowDefinition:
        """Create a new automation workflow"""

        workflow_id = str(uuid.uuid4())

        workflow = WorkflowDefinition(
            workflow_id=workflow_id,
            name=workflow_definition['name'],
            description=workflow_definition.get('description', ''),
            tenant_id=self.tenant_id,
            automation_type=AutomationType(workflow_definition['automation_type']),
            trigger_conditions=workflow_definition.get('trigger_conditions', {}),
            processing_steps=workflow_definition.get('processing_steps', []),
            output_configuration=workflow_definition.get('output_configuration', {}),
            status=WorkflowStatus.DRAFT,
            created_by=workflow_definition.get('created_by', 'system'),
            created_at=datetime.now(),
            last_modified=datetime.now()
        )

        # Store in database
        await self._store_workflow_definition(workflow)

        return workflow

    async def execute_workflow(self, workflow_id: str, input_data: Dict[str, Any]) -> WorkflowExecution:
        """Execute an automation workflow"""

        execution_id = str(uuid.uuid4())
        start_time = datetime.now()

        execution = WorkflowExecution(
            execution_id=execution_id,
            workflow_id=workflow_id,
            tenant_id=self.tenant_id,
            status=ProcessingStatus.PROCESSING,
            input_data=input_data,
            output_data=None,
            started_at=start_time,
            completed_at=None,
            error_message=None,
            processing_logs=[]
        )

        try:
            # Get workflow definition
            workflow = await self._get_workflow_definition(workflow_id)
            if not workflow:
                raise ValueError(f"Workflow {workflow_id} not found")

            # Store execution record
            await self._store_workflow_execution(execution)

            # Execute workflow based on type
            processor = self.workflow_processors.get(workflow.automation_type)
            if not processor:
                raise ValueError(f"No processor for automation type: {workflow.automation_type}")

            output_data = await processor(workflow, input_data, execution)

            # Update execution with results
            execution.status = ProcessingStatus.COMPLETED
            execution.output_data = output_data
            execution.completed_at = datetime.now()
            execution.execution_time_seconds = (execution.completed_at - start_time).total_seconds()

            await self._update_workflow_execution(execution)

            return execution

        except Exception as e:
            logger.error(f"Workflow execution failed: {str(e)}")

            execution.status = ProcessingStatus.FAILED
            execution.error_message = str(e)
            execution.completed_at = datetime.now()
            execution.execution_time_seconds = (execution.completed_at - start_time).total_seconds()

            await self._update_workflow_execution(execution)

            return execution

    async def _execute_document_extraction(self, workflow: WorkflowDefinition,
                                         input_data: Dict, execution: WorkflowExecution) -> Dict[str, Any]:
        """Execute document extraction workflow"""

        document_ids = input_data.get('document_ids', [])
        extraction_results = []

        for document_id in document_ids:
            result = await self.process_document(document_id, workflow.processing_steps[0] if workflow.processing_steps else {})
            extraction_results.append(result)

        return {
            'extracted_documents': len(extraction_results),
            'successful_extractions': len([r for r in extraction_results if r.status == ProcessingStatus.COMPLETED]),
            'results': [asdict(r) for r in extraction_results]
        }

    async def _execute_data_transformation(self, workflow: WorkflowDefinition,
                                         input_data: Dict, execution: WorkflowExecution) -> Dict[str, Any]:
        """Execute data transformation workflow"""

        # Placeholder for data transformation logic
        transformation_rules = workflow.processing_steps
        source_data = input_data.get('data', {})

        transformed_data = source_data.copy()

        # Apply transformation rules
        for rule in transformation_rules:
            rule_type = rule.get('type')
            if rule_type == 'field_mapping':
                # Map fields according to rules
                field_mappings = rule.get('mappings', {})
                for old_field, new_field in field_mappings.items():
                    if old_field in transformed_data:
                        transformed_data[new_field] = transformed_data.pop(old_field)

            elif rule_type == 'data_validation':
                # Validate data according to rules
                validation_rules = rule.get('rules', [])
                validation_results = []
                for validation_rule in validation_rules:
                    # Implement validation logic
                    validation_results.append({'rule': validation_rule, 'passed': True})
                transformed_data['validation_results'] = validation_results

        return {
            'transformation_applied': True,
            'rules_processed': len(transformation_rules),
            'transformed_data': transformed_data
        }

    # Helper methods for database operations
    async def _store_document_metadata(self, metadata: DocumentMetadata):
        """Store document metadata in database"""

        async with self.db_pool.acquire() as conn:
            await conn.execute("""
                INSERT INTO operations.documents (
                    document_id, filename, document_type, file_size, tenant_id,
                    user_id, checksum, mime_type, processing_status, metadata, tags
                ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11)
            """,
                metadata.document_id, metadata.filename, metadata.document_type.value,
                metadata.file_size, metadata.tenant_id, metadata.user_id,
                metadata.checksum, metadata.mime_type, metadata.processing_status.value,
                json.dumps(metadata.metadata), json.dumps(metadata.tags)
            )

    async def _get_document_metadata(self, document_id: str) -> Optional[DocumentMetadata]:
        """Get document metadata from database"""

        async with self.db_pool.acquire() as conn:
            row = await conn.fetchrow("""
                SELECT * FROM operations.documents WHERE document_id = $1 AND tenant_id = $2
            """, document_id, self.tenant_id)

            if row:
                return DocumentMetadata(
                    document_id=row['document_id'],
                    filename=row['filename'],
                    document_type=DocumentType(row['document_type']),
                    file_size=row['file_size'],
                    upload_timestamp=row['created_at'],
                    tenant_id=row['tenant_id'],
                    user_id=row['user_id'],
                    checksum=row['checksum'],
                    mime_type=row['mime_type'],
                    processing_status=ProcessingStatus(row['processing_status']),
                    metadata=json.loads(row['metadata']) if row['metadata'] else {},
                    tags=json.loads(row['tags']) if row['tags'] else []
                )

        return None

    def _get_mime_type(self, filename: str) -> str:
        """Get MIME type from filename"""

        extension = Path(filename).suffix.lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.json': 'application/json'
        }

        return mime_types.get(extension, 'application/octet-stream')

    # Additional helper methods would be implemented here...
    async def _queue_document_processing(self, document_id: str):
        """Queue document for processing"""
        pass

    async def _update_processing_status(self, document_id: str, status: ProcessingStatus):
        """Update document processing status"""
        pass

    async def _store_processing_result(self, result: ProcessingResult):
        """Store processing result"""
        pass

# Database schema for operations
# Database schema for operations
OPERATIONS_SCHEMA_SQL = """
-- Operations automation schema
CREATE SCHEMA IF NOT EXISTS operations;

-- Documents table
CREATE TABLE IF NOT EXISTS operations.documents (
    document_id UUID PRIMARY KEY,
    filename VARCHAR(500) NOT NULL,
    document_type VARCHAR(50) NOT NULL,
    file_size BIGINT NOT NULL,
    tenant_id UUID REFERENCES tenant_management.tenants(id) ON DELETE CASCADE,
    user_id UUID NOT NULL,
    checksum VARCHAR(128) NOT NULL,
    mime_type VARCHAR(200) NOT NULL,
    processing_status VARCHAR(50) DEFAULT 'pending',
    extracted_text TEXT,
    metadata JSONB DEFAULT '{}',
    tags JSONB DEFAULT '[]',
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Processing results table
CREATE TABLE IF NOT EXISTS operations.processing_results (
    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    document_id UUID REFERENCES operations.documents(document_id) ON DELETE CASCADE,
    processing_type VARCHAR(100) NOT NULL,
    status VARCHAR(50) NOT NULL,
    extracted_data JSONB DEFAULT '{}',
    confidence_score FLOAT DEFAULT 0,
    processing_time_seconds FLOAT DEFAULT 0,
    error_message TEXT,
    suggestions JSONB DEFAULT '[]',
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Workflow definitions table
CREATE TABLE IF NOT EXISTS operations.workflow_definitions (
    workflow_id UUID PRIMARY KEY,
    name VARCHAR(500) NOT NULL,
    description TEXT,
    tenant_id UUID REFERENCES tenant_management.tenants(id) ON DELETE CASCADE,
    automation_type VARCHAR(100) NOT NULL,
    trigger_conditions JSONB DEFAULT '{}',
    processing_steps JSONB DEFAULT '[]',
    output_configuration JSONB DEFAULT '{}',
    status VARCHAR(50) DEFAULT 'draft',
    created_by UUID NOT NULL,
    execution_count INTEGER DEFAULT 0,
    success_rate FLOAT DEFAULT 0,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    last_modified TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Workflow executions table
CREATE TABLE IF NOT EXISTS operations.workflow_executions (
    execution_id UUID PRIMARY KEY,
    workflow_id UUID REFERENCES operations.workflow_definitions(workflow_id) ON DELETE CASCADE,
    tenant_id UUID REFERENCES tenant_management.tenants(id) ON DELETE CASCADE,
    status VARCHAR(50) DEFAULT 'processing',
    input_data JSONB DEFAULT '{}',
    output_data JSONB DEFAULT '{}',
    error_message TEXT,
    processing_logs JSONB DEFAULT '[]',
    execution_time_seconds FLOAT,
    started_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    completed_at TIMESTAMP WITH TIME ZONE
);

-- Document processing queue
CREATE TABLE IF NOT EXISTS operations.processing_queue (
    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    document_id UUID REFERENCES operations.documents(document_id) ON DELETE CASCADE,
    tenant_id UUID REFERENCES tenant_management.tenants(id) ON DELETE CASCADE,
    priority INTEGER DEFAULT 5,
    status VARCHAR(50) DEFAULT 'queued',
    processing_options JSONB DEFAULT '{}',
    retry_count INTEGER DEFAULT 0,
    max_retries INTEGER DEFAULT 3,
    queued_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    started_at TIMESTAMP WITH TIME ZONE,
    completed_at TIMESTAMP WITH TIME ZONE
);

-- Automation triggers table
CREATE TABLE IF NOT EXISTS operations.automation_triggers (
    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    trigger_name VARCHAR(200) NOT NULL,
    tenant_id UUID REFERENCES tenant_management.tenants(id) ON DELETE CASCADE,
    trigger_type VARCHAR(100) NOT NULL,
    conditions JSONB DEFAULT '{}',
    workflow_id UUID REFERENCES operations.workflow_definitions(workflow_id) ON DELETE CASCADE,
    is_active BOOLEAN DEFAULT true,
    execution_count INTEGER DEFAULT 0,
    last_triggered TIMESTAMP WITH TIME ZONE,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Indexes for performance
CREATE INDEX IF NOT EXISTS idx_documents_tenant_status
ON operations.documents(tenant_id, processing_status);

CREATE INDEX IF NOT EXISTS idx_documents_type_created
ON operations.documents(document_type, created_at DESC);

CREATE INDEX IF NOT EXISTS idx_processing_results_document
ON operations.processing_results(document_id, created_at DESC);

CREATE INDEX IF NOT EXISTS idx_workflow_executions_tenant_status
ON operations.workflow_executions(tenant_id, status);

CREATE INDEX IF NOT EXISTS idx_processing_queue_status_priority
ON operations.processing_queue(status, priority DESC, queued_at ASC);

CREATE INDEX IF NOT EXISTS idx_automation_triggers_active
ON operations.automation_triggers(tenant_id, is_active, trigger_type);

-- Views for common queries
CREATE OR REPLACE VIEW operations.document_processing_summary AS
SELECT
    d.tenant_id,
    d.document_type,
    COUNT(*) as total_documents,
    SUM(CASE WHEN d.processing_status = 'completed' THEN 1 ELSE 0 END) as completed,
    SUM(CASE WHEN d.processing_status = 'failed' THEN 1 ELSE 0 END) as failed,
    SUM(CASE WHEN d.processing_status = 'processing' THEN 1 ELSE 0 END) as processing,
    AVG(pr.confidence_score) as avg_confidence,
    AVG(pr.processing_time_seconds) as avg_processing_time
FROM operations.documents d
LEFT JOIN operations.processing_results pr ON d.document_id = pr.document_id
GROUP BY d.tenant_id, d.document_type;

-- Functions for automation
CREATE OR REPLACE FUNCTION operations.trigger_workflow_execution()
RETURNS TRIGGER AS $$
BEGIN
    -- Check for matching automation triggers
    INSERT INTO operations.workflow_executions (
        execution_id, workflow_id, tenant_id, input_data
    )
    SELECT
        gen_random_uuid(),
        at.workflow_id,
        NEW.tenant_id,
        json_build_object('document_id', NEW.document_id)
    FROM operations.automation_triggers at
    WHERE at.tenant_id = NEW.tenant_id
    AND at.is_active = true
    AND at.trigger_type = 'document_upload'
    AND (at.conditions IS NULL OR
         (at.conditions->>'document_type' IS NULL OR
          at.conditions->>'document_type' = NEW.document_type));

    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

-- Trigger to auto-execute workflows on document upload
CREATE TRIGGER trigger_document_workflow_execution
    AFTER INSERT ON operations.documents
    FOR EACH ROW
    EXECUTE FUNCTION operations.trigger_workflow_execution();

-- Comments for documentation
COMMENT ON TABLE operations.documents IS 'Stores document metadata and processing status';
COMMENT ON TABLE operations.processing_results IS 'Stores results of document processing operations';
COMMENT ON TABLE operations.workflow_definitions IS 'Defines automation workflows and their configuration';
COMMENT ON TABLE operations.workflow_executions IS 'Tracks execution of automation workflows';
COMMENT ON TABLE operations.processing_queue IS 'Queue for managing document processing jobs';
COMMENT ON TABLE operations.automation_triggers IS 'Defines triggers that automatically execute workflows';
"""

# Pydantic models for API
class DocumentUploadRequest(BaseModel):
    filename: str = Field(..., description="Name of the file")
    file_content: str = Field(..., description="Base64 encoded file content")
    metadata: Optional[Dict[str, Any]] = Field(None, description="Additional metadata")
    processing_options: Optional[Dict[str, Any]] = Field(None, description="Processing options")

class DocumentProcessingRequest(BaseModel):
    document_id: str = Field(..., description="Document ID to process")
    processing_options: Optional[Dict[str, Any]] = Field(None, description="Processing options")

class WorkflowCreateRequest(BaseModel):
    name: str = Field(..., description="Workflow name")
    description: Optional[str] = Field(None, description="Workflow description")
    automation_type: str = Field(..., description="Type of automation")
    trigger_conditions: Optional[Dict[str, Any]] = Field(None, description="Trigger conditions")
    processing_steps: List[Dict[str, Any]] = Field(..., description="Processing steps")
    output_configuration: Optional[Dict[str, Any]] = Field(None, description="Output configuration")

class WorkflowExecuteRequest(BaseModel):
    workflow_id: str = Field(..., description="Workflow ID to execute")
    input_data: Dict[str, Any] = Field(..., description="Input data for execution")

# Export main classes
__all__ = [
    'OperationsAgent', 'DocumentMetadata', 'ProcessingResult', 'WorkflowDefinition', 'WorkflowExecution',
    'DocumentType', 'ProcessingStatus', 'WorkflowStatus', 'AutomationType',
    'DocumentUploadRequest', 'DocumentProcessingRequest', 'WorkflowCreateRequest', 'WorkflowExecuteRequest',
    'OPERATIONS_SCHEMA_SQL'
]
